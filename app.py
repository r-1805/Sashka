import base64
import json
import os
from datetime import datetime, timedelta
import logging
from io import StringIO, BytesIO
import io
from time import sleep
from werkzeug.utils import secure_filename
from flask import Flask, render_template, url_for, Response, request, jsonify, redirect, send_file, make_response, flash, abort
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, SubmitField
from wtforms.validators import InputRequired, Length, ValidationError
from werkzeug.security import generate_password_hash, check_password_hash
from sqlalchemy import create_engine, Column, String, Integer, DateTime, ForeignKey, Text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import relationship, sessionmaker
from flask_cors import CORS
from AI import recognize_question
import subprocess
import sys
import secrets
import re
import glob
import random
import markdown
from docx import Document
from docx2pdf import convert as docx2pdf_convert
import tempfile

from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

app = Flask(__name__, template_folder="templates")

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///intents.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
CORS(app)
db_flask = SQLAlchemy(app)

# Генерация и установка секретного ключа
app.config['SECRET_KEY'] = secrets.token_hex(16)

# Настройка Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Модель пользователя
class User(db_flask.Model, UserMixin):
    __tablename__ = 'users'
    id = db_flask.Column(db_flask.Integer, primary_key=True)
    username = db_flask.Column(db_flask.String(150), unique=True, nullable=False)
    password_hash = db_flask.Column(db_flask.String(150), nullable=False)
    histories = db_flask.relationship('History', backref='user', lazy=True)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

# Модель для хранения истории генерации статей
class History(db_flask.Model):
    __tablename__ = 'history'
    id = db_flask.Column(db_flask.Integer, primary_key=True)
    user_id = db_flask.Column(db_flask.Integer, db_flask.ForeignKey('users.id'), nullable=False)
    timestamp = db_flask.Column(db_flask.DateTime, default=datetime.utcnow)
    raw_text = db_flask.Column(db_flask.Text, nullable=False)
    word_count = db_flask.Column(db_flask.Integer, nullable=False)
    style = db_flask.Column(db_flask.String(50), nullable=False)
    audience = db_flask.Column(db_flask.String(50), nullable=False)
    references_format = db_flask.Column(db_flask.String(50), nullable=False)
    generated_text = db_flask.Column(db_flask.Text, nullable=False)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@app.route('/')
def index():
    return redirect(url_for('generate'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        confirm_password = request.form['confirm_password']

        if password != confirm_password:
            flash('Пароли не совпадают', 'danger')
            return redirect(url_for('register'))

        if User.query.filter_by(username=username).first():
            flash('Этот логин уже используется', 'danger')
            return redirect(url_for('register'))

        new_user = User(username=username)
        new_user.set_password(password)

        db_flask.session.add(new_user)
        db_flask.session.commit()

        flash('Регистрация прошла успешно! Теперь вы можете войти.', 'success')
        return redirect(url_for('login'))

    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password):
            login_user(user)
            flash('Вы успешно вошли!', 'success')
            logging.info(f'Пользователь {username} вошел в систему')
            return redirect(url_for('generate'))
        else:
            flash('Неверный логин или пароль', 'danger')
            logging.warning(f'Неудачная попытка входа для пользователя {username}')
            return redirect(url_for('login'))

    return render_template('login.html')

@app.route('/generate', methods=['GET'])
@login_required
def generate():
    return render_template('generate.html')

def add_hyperlink(paragraph, text, url):
    """
    Добавляет гиперссылку в параграф.
    Ссылка будет отображаться синим цветом и подчеркнута.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def convert_markdown_to_word(markdown_text):
    """
    Преобразует markdown-текст в DOCX с сохранением форматирования.
    Обрабатываются:
      - Заголовки (например, "### Заголовок")
      - Нумерованные списки (строки, начинающиеся с "1. ")
      - Жирный (**text** или __text__)
      - Курсив (*text* или _text_)
      - Гиперссылки ([текст](url))
    При обработке заголовков для каждого уровня устанавливается соответствующий размер шрифта.
    """
    document = Document()
    lines = markdown_text.splitlines()
    token_pattern = re.compile(r'(\[[^\]]+\]\([^)]+\)|\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)')
    hyperlink_pattern = re.compile(r'^\[([^\]]+)\]\(([^)]+)\)$')
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            document.add_paragraph("")
            continue
        # Обработка заголовков (начинаются с #)
        if stripped_line.startswith("#"):
            m = re.match(r'^(#+)\s*(.*)$', stripped_line)
            if m:
                level = len(m.group(1))
                text = m.group(2)
                heading = document.add_heading(text, level=level)
                # Устанавливаем размер шрифта для заголовка в зависимости от уровня
                size_map = {1: Pt(24), 2: Pt(20), 3: Pt(16), 4: Pt(14), 5: Pt(12), 6: Pt(10)}
                for run in heading.runs:
                    run.font.size = size_map.get(level, Pt(12))
                continue
        # Обработка нумерованных списков
        if re.match(r'^\d+\.\s', stripped_line):
            paragraph = document.add_paragraph(stripped_line, style='List Number')
            continue
        # Остальные строки — обычный текст с форматированием
        paragraph = document.add_paragraph()
        tokens = token_pattern.split(line)
        tokens = [t for t in tokens if t]  # фильтруем пустые токены
        for token in tokens:
            hyperlink_match = hyperlink_pattern.match(token)
            if hyperlink_match:
                link_text = hyperlink_match.group(1)
                url = hyperlink_match.group(2)
                add_hyperlink(paragraph, link_text, url)
            elif token.startswith('**') and token.endswith('**'):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('__') and token.endswith('__'):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('*') and token.endswith('*'):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            elif token.startswith('_') and token.endswith('_'):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                paragraph.add_run(token)
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def convert_word_to_pdf(word_stream):
    """
    Конвертирует DOCX (в виде потока) в PDF с использованием docx2pdf.
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
        tmp_docx.write(word_stream.read())
        tmp_docx_path = tmp_docx.name

    pdf_path = tmp_docx_path.replace('.docx', '.pdf')
    try:
        docx2pdf_convert(tmp_docx_path, pdf_path)
        with open(pdf_path, 'rb') as f:
            pdf_data = f.read()
    finally:
        os.remove(tmp_docx_path)
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
    return pdf_data

@app.route('/generate_article', methods=['POST'])
@login_required
def generate_article():
    raw_text = request.form['raw_text']
    word_count = request.form['word_count']
    style = request.form['style']
    audience = request.form['audience']
    references_format = request.form['references_format']
    journal_link = request.form.get('journal_link', '').strip()

    generated_text = recognize_question(raw_text, word_count, style, audience, references_format)
    if journal_link:
        # Добавляем ссылку в markdown-формате
        generated_text += f"\n\n[Ссылка на журнал]({journal_link})"

    # Преобразуем markdown в HTML для отображения на странице
    generated_text_html = markdown.markdown(generated_text)

    history_entry = History(
        user_id=current_user.id,
        raw_text=raw_text,
        word_count=word_count,
        style=style,
        audience=audience,
        references_format=references_format,
        generated_text=generated_text,  # сохраняем оригинальный markdown
        timestamp=datetime.utcnow()
    )
    db_flask.session.add(history_entry)
    db_flask.session.commit()

    return render_template('generate.html', generated_text=generated_text_html)

@app.route('/history')
@login_required
def history():
    user_history = History.query.filter_by(user_id=current_user.id).order_by(History.timestamp.desc()).all()
    return render_template('history.html', history=user_history)

@app.route('/download_word/<int:entry_id>')
@login_required
def download_word(entry_id):
    entry = History.query.filter_by(id=entry_id, user_id=current_user.id).first_or_404()
    word_stream = convert_markdown_to_word(entry.generated_text)
    filename = f"article_{entry.id}.docx"
    return send_file(word_stream, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/download_pdf/<int:entry_id>')
@login_required
def download_pdf(entry_id):
    entry = History.query.filter_by(id=entry_id, user_id=current_user.id).first_or_404()
    word_stream = convert_markdown_to_word(entry.generated_text)
    pdf_data = convert_word_to_pdf(word_stream)
    filename = f"article_{entry.id}.pdf"
    return send_file(io.BytesIO(pdf_data), as_attachment=True, download_name=filename,
                     mimetype='application/pdf')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Вы успешно вышли из системы', 'info')
    return redirect(url_for('login'))

with app.app_context():
    db_flask.create_all()

if __name__ == "__main__":
    app.run(debug=True)
